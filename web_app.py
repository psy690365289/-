import streamlit as st
import pandas as pd
import io
import base64
from docx import Document
from docx.shared import Pt
import re
import os


# ================== 辅助函数 ==================
def format_number(value):
    try:
        num = round(float(value), 2)
        is_negative = num < 0
        num = abs(num)
        formatted = "{:.2f}".format(num).rstrip('0').rstrip('.')

        if '.' in formatted:
            integer_part, decimal_part = formatted.split('.')
        else:
            integer_part, decimal_part = formatted, ''

        integer_with_commas = ''
        for i, ch in enumerate(reversed(integer_part)):
            if i > 0 and i % 3 == 0:
                integer_with_commas = ',' + integer_with_commas
            integer_with_commas = ch + integer_with_commas

        result = integer_with_commas
        if decimal_part:
            result = f"{result}.{decimal_part}"
        if is_negative:
            result = '-' + result

        return result
    except:
        return str(value)


def clean_feature_line(line):
    line = line.strip()
    if not line:
        return ""
    cleaned = re.sub(r'^\d+[\.、]\s*', '', line)
    return cleaned if cleaned else line


def parse_item_name_and_features(text):
    if not isinstance(text, str):
        return str(text)

    text = text.strip()
    if not text:
        return ""

    lines = [line.strip() for line in text.split('\n') if line.strip()]

    if not lines:
        return text

    project_name = lines[0]
    features = []

    if len(lines) > 1:
        for line in lines[1:]:
            cleaned_feature = clean_feature_line(line)
            if cleaned_feature:
                features.append(cleaned_feature)
    else:
        single_line = lines[0]
        match = re.search(r'\s+\d+[\.、]\s*', project_name)
        if match:
            name_part = project_name[:match.start()].strip()
            features_part = project_name[match.start():].strip()
            project_name = name_part
            feature_sections = re.split(r'\s+(?=\d+[\.、])', features_part)
            for section in feature_sections:
                cleaned = clean_feature_line(section)
                if cleaned:
                    features.append(cleaned)
        else:
            pattern = r'(?<!\d)(\d+[\.、]\s*[^。，；!?]+(?:[。，；!?](?!\s*\d+[\.、])[^。，；!?]*)*)'
            matches = re.findall(pattern, text)
            if matches:
                for match in matches:
                    cleaned = clean_feature_line(match)
                    if cleaned and cleaned not in project_name:
                        features.append(cleaned)
                for feature in features:
                    project_name = project_name.replace(feature, '').strip()

    if features:
        unique_features = []
        for feature in features:
            feature = feature.strip()
            feature = re.sub(r'[，。；!?]+$', '', feature)
            if feature and feature not in unique_features:
                unique_features.append(feature)
        if unique_features:
            return f"{project_name}（{'、'.join(unique_features)}）"

    return project_name


def read_excel_data(file_content, reduction_range=None, increase_range=None):
    try:
        df = pd.read_excel(
            io.BytesIO(file_content),
            sheet_name="【分部1】分部分项清单对比表",
            skiprows=3,
            header=None,
            usecols=[1, 2, 4, 5, 6, 8, 10, 12]
        )

        df.columns = ["项目编码", "项目名称与特征", "计量单位",
                      "送审工程量", "送审单价", "审定工程量", "审定单价", "审减金额"]

        df = df.dropna(subset=["项目编码", "项目名称与特征"])

        numeric_cols = ["送审工程量", "审定工程量", "送审单价", "审定单价", "审减金额"]
        for col in numeric_cols:
            df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)

        df['项目名称与特征'] = df['项目名称与特征'].astype(str).str.strip()
        df['项目名称与特征'] = df['项目名称与特征'].str.replace(r'[ \t]+', ' ', regex=True)
        df['项目名称与特征'] = df['项目名称与特征'].apply(parse_item_name_and_features)

        conditions = []
        if reduction_range:
            min_reduction, max_reduction = reduction_range
            conditions.append((df['审减金额'] <= -min_reduction) & (df['审减金额'] >= -max_reduction))
        if increase_range:
            min_increase, max_increase = increase_range
            conditions.append((df['审减金额'] >= min_increase) & (df['审减金额'] <= max_increase))

        if conditions:
            combined_condition = pd.concat(conditions, axis=1).any(axis=1)
            df = df[combined_condition]

        df = df[df['审减金额'] != 0]
        return df
    except Exception as e:
        raise ValueError(f"Excel处理失败: {str(e)}")


def generate_word_report(df):
    if df.empty:
        raise ValueError("没有符合条件的数据！")

    try:
        doc = Document()
        doc.add_heading('工程审核报告（筛选版）', level=1)

        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(10.5)

        for i, row in enumerate(df.itertuples(), start=1):
            name = str(row.项目名称与特征).strip()
            unit = str(row.计量单位).strip()
            qty_sent = row.送审工程量
            qty_audited = row.审定工程量

            if round(qty_sent, 2) == round(qty_audited, 2):
                qty_text = f"工程量为{format_number(qty_sent)}{unit}，"
            else:
                qty_text = f"送审工程量为{format_number(qty_sent)}{unit}，"

            text_parts = [f"{i}、{name}，", qty_text]

            if round(qty_sent, 2) != round(qty_audited, 2):
                text_parts.append(f"审核工程量为{format_number(qty_audited)}{unit}，")

            amount = row.审减金额
            amount_type = "审增" if amount > 0 else "审减"
            formatted_amount = format_number(abs(amount))

            text_parts.extend([
                f"送审单价为{format_number(row.送审单价)}元/{unit}，",
                f"审核单价为{format_number(row.审定单价)}元/{unit}，",
                f"此项{amount_type}金额为{formatted_amount}元。"
            ])
            doc.add_paragraph(''.join(text_parts))

        doc.add_heading('汇总统计', level=2)
        increase_df = df[df['审减金额'] > 0]
        reduction_df = df[df['审减金额'] < 0]

        total_increase = increase_df['审减金额'].sum()
        total_reduction = reduction_df['审减金额'].abs().sum()

        def get_range(data, is_increase=True):
            if data.empty:
                return "无"
            values = data if is_increase else data.abs()
            return f"{format_number(values.min())}元 ~ {format_number(values.max())}元"

        summary_text = (
            f"▪ 总审增金额：{format_number(total_increase) if total_increase != 0 else '无'}\n"
            f"▪ 总审减金额：{format_number(total_reduction) if total_reduction != 0 else '无'}\n"
            f"▪ 筛选记录数：{len(df)}条\n"
            f"▪ 审增区间：{get_range(increase_df['审减金额'])}\n"
            f"▪ 审减区间：{get_range(reduction_df['审减金额'], False)}"
        )
        doc.add_paragraph(summary_text)

        # 保存到字节流
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        raise ValueError(f"Word生成失败: {str(e)}")


# ================== 网站主程序 ==================
st.set_page_config(page_title="工程审核报告生成器", page_icon="📊", layout="wide")

st.title("📊 工程审核报告生成器（在线版）")
st.markdown("---")

# 统计使用人数（简单版，记录session）
if 'usage_count' not in st.session_state:
    st.session_state.usage_count = 0

st.session_state.usage_count += 1

# 显示使用人数
st.sidebar.info(f"👥 今日使用人数：{st.session_state.usage_count}")

# 文件上传区
uploaded_file = st.file_uploader("📤 上传Excel文件", type=['xlsx'])

# 筛选条件
st.subheader("🔍 筛选条件")
col1, col2 = st.columns(2)

with col1:
    st.markdown("**审减金额范围**")
    reduction_frame = ttk.LabelFrame(main_frame, text="审减金额范围")
    reduction_frame.pack(pady=5, fill="x")
    ttk.Label(reduction_frame, text="最小 ≥").grid(row=0, column=0, padx=5)
    self.reduction_min = ttk.Entry(reduction_frame, width=10)
    self.reduction_min.grid(row=0, column=1, padx=5)
    ttk.Label(reduction_frame, text="最大 ≤").grid(row=0, column=2, padx=5)
    self.reduction_max = ttk.Entry(reduction_frame, width=10)
    self.reduction_max.grid(row=0, column=3, padx=5)

with col2:
    st.markdown("**审增金额范围**")
    increase_frame = ttk.LabelFrame(main_frame, text="审增金额范围")
    increase_frame.pack(pady=5, fill="x")
    ttk.Label(increase_frame, text="最小 ≥").grid(row=0, column=0, padx=5)
    self.increase_min = ttk.Entry(increase_frame, width=10)
    self.increase_min.grid(row=0, column=1, padx=5)
    ttk.Label(increase_frame, text="最大 ≤").grid(row=0, column=2, padx=5)
    self.increase_max = ttk.Entry(increase_frame, width=10)
    self.increase_max.grid(row=0, column=3, padx=5)

st.markdown("---")

# 处理按钮
if st.button("🚀 开始生成报告", type="primary"):
    if uploaded_file is not None:
        with st.spinner("正在处理中..."):
            # 设置筛选范围
            reduction_range = (reduction_min, reduction_max) if (reduction_min > 0 or reduction_max > 0) else None
            increase_range = (increase_min, increase_max) if (increase_min > 0 or increase_max > 0) else None

            try:
                # 读取Excel数据
                excel_content = uploaded_file.read()
                df = read_excel_data(excel_content, reduction_range, increase_range)

                if df.empty:
                    st.warning("⚠️ 没有找到符合条件的数据！")
                else:
                    # 显示结果预览
                    st.success(f"✅ 找到 {len(df)} 条记录")

                    # 显示数据预览
                    with st.expander("📋 查看数据预览", expanded=False):
                        st.dataframe(df[['项目名称与特征', '审减金额']].head(10))

                    # 生成Word报告
                    word_buffer = generate_word_report(df)

                    # 下载按钮
                    st.download_button(
                        label="📥 下载Word报告",
                        data=word_buffer,
                        file_name="工程审核报告.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                    # 显示统计信息
                    increase_total = df[df['审减金额'] > 0]['审减金额'].sum()
                    reduction_total = df[df['审减金额'] < 0]['审减金额'].abs().sum()

                    st.info(f"""
                    **📊 统计结果：**
                    - 总审增金额：{format_number(increase_total)}元
                    - 总审减金额：{format_number(reduction_total)}元
                    """)

            except Exception as e:
                st.error(f"❌ 处理失败：{str(e)}")
    else:
        st.warning("⚠️ 请先上传Excel文件！")

# 管理员查看区（密码保护）
st.sidebar.markdown("---")
st.sidebar.subheader("👑 管理员面板")

admin_password = st.sidebar.text_input("管理员密码", type="password")
if admin_password == "admin123":  # 你可以改密码
    st.sidebar.success("✅ 管理员登录成功")

    # 显示所有上传记录（简化版）
    if 'uploaded_files' not in st.session_state:
        st.session_state.uploaded_files = []

    if uploaded_file:
        st.session_state.uploaded_files.append({
            'name': uploaded_file.name,
            'time': pd.Timestamp.now(),
            'size': f"{uploaded_file.size / 1024:.1f} KB"
        })

    if st.session_state.uploaded_files:
        st.sidebar.subheader("📁 最近上传的文件")
        for file_info in list(reversed(st.session_state.uploaded_files))[:5]:
            st.sidebar.text(f"📄 {file_info['name']}")
    else:
        st.sidebar.text("暂无上传记录")

st.markdown("---")

st.caption("💡 提示：请确保Excel文件包含名为【分部1】分部分项清单对比表的工作表")
